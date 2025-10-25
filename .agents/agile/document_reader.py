#!/usr/bin/env python3
"""
Document Reader - Multi-Format Requirements Document Reader

Reads requirements from various file formats:
- PDF (.pdf)
- Microsoft Word (.docx, .doc)
- Microsoft Excel (.xlsx, .xls)
- LibreOffice Writer (.odt)
- LibreOffice Calc (.ods)
- Plain Text (.txt)
- Markdown (.md)
- CSV (.csv)
- HTML (.html)
- Jupyter Notebooks (.ipynb)
"""

import os
from pathlib import Path
from typing import Optional, Dict, List
import csv

from artemis_exceptions import (
    UnsupportedDocumentFormatError,
    DocumentReadError,
    wrap_exception
)
from jupyter_notebook_handler import JupyterNotebookReader


class DocumentReader:
    """
    Read requirements documents from various formats

    Handles multiple file formats and extracts plain text for LLM processing.
    """

    # Performance: Dict dispatch for O(1) extension lookup instead of O(n) if/elif chain
    EXTENSION_HANDLERS = {
        '.pdf': '_read_pdf',
        '.docx': '_read_word',
        '.doc': '_read_word',
        '.xlsx': '_read_excel',
        '.xls': '_read_excel',
        '.odt': '_read_odt',
        '.ods': '_read_ods',
        '.txt': '_read_text',
        '.md': '_read_text',
        '.markdown': '_read_text',
        '.csv': '_read_csv',
        '.html': '_read_html',
        '.htm': '_read_html',
        '.ipynb': '_read_ipynb'
    }

    def __init__(self, verbose: bool = True):
        """
        Initialize Document Reader

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which document libraries are available"""
        self.has_pdf = False
        self.has_docx = False
        self.has_openpyxl = False
        self.has_odf = False
        self.has_pypandoc = False

        # Check for PyPDF2 or pdfplumber
        try:
            import PyPDF2
            self.has_pdf = True
            self.pdf_library = "PyPDF2"
        except ImportError:
            try:
                import pdfplumber
                self.has_pdf = True
                self.pdf_library = "pdfplumber"
            except ImportError:
                self.log("⚠️  PDF support not available. Install: pip install PyPDF2 or pdfplumber")

        # Check for python-docx (Word documents)
        try:
            import docx
            self.has_docx = True
        except ImportError:
            self.log("⚠️  Word document support not available. Install: pip install python-docx")

        # Check for openpyxl (Excel)
        try:
            import openpyxl
            self.has_openpyxl = True
        except ImportError:
            self.log("⚠️  Excel support not available. Install: pip install openpyxl")

        # Check for odfpy (LibreOffice)
        try:
            from odf import text, teletype
            from odf.opendocument import load
            self.has_odf = True
        except ImportError:
            self.log("⚠️  LibreOffice support not available. Install: pip install odfpy")

        # Check for pypandoc (universal converter)
        try:
            import pypandoc
            self.has_pypandoc = True
        except ImportError:
            pass

    def read_document(self, file_path: str) -> str:
        """
        Read document and extract text content

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        try:
            path = Path(file_path)

            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            extension = path.suffix.lower()

            self.log(f"📄 Reading {extension} file: {path.name}")

            # Performance: Use dict dispatch for O(1) lookup
            handler_name = self.EXTENSION_HANDLERS.get(extension)
            if handler_name:
                handler = getattr(self, handler_name)
                return handler(file_path)
            else:
                # Try pypandoc as fallback
                if self.has_pypandoc:
                    return self._read_with_pandoc(file_path)
                else:
                    raise UnsupportedDocumentFormatError(
                        f"Unsupported file format: {extension}",
                        context={"file_path": file_path, "extension": extension}
                    )

        except FileNotFoundError:
            raise  # Re-raise as-is
        except UnsupportedDocumentFormatError:
            raise  # Re-raise as-is
        except RuntimeError as e:
            # Library not available errors
            raise DocumentReadError(
                str(e),
                context={"file_path": file_path, "extension": Path(file_path).suffix},
                original_exception=e
            )
        except Exception as e:
            raise wrap_exception(
                e,
                DocumentReadError,
                f"Error reading document: {file_path}",
                context={"file_path": file_path}
            )

    def _read_pdf(self, file_path: str) -> str:
        """Read PDF file"""
        if not self.has_pdf:
            raise RuntimeError("PDF support not available. Install: pip install PyPDF2 or pdfplumber")

        if self.pdf_library == "PyPDF2":
            import PyPDF2
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
            return "\n".join(text_content)

        else:  # pdfplumber
            import pdfplumber
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            return "\n".join(text_content)

    def _read_word(self, file_path: str) -> str:
        """Read Microsoft Word (.docx) file"""
        if not self.has_docx:
            raise RuntimeError("Word document support not available. Install: pip install python-docx")

        import docx

        doc = docx.Document(file_path)
        text_content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_content.append(row_text)

        return "\n".join(text_content)

    def _read_excel(self, file_path: str) -> str:
        """Read Microsoft Excel (.xlsx, .xls) file"""
        if not self.has_openpyxl:
            raise RuntimeError("Excel support not available. Install: pip install openpyxl")

        import openpyxl

        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_content.append(f"\n=== Sheet: {sheet_name} ===\n")

            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to strings
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    text_content.append(" | ".join(row_values))

        return "\n".join(text_content)

    def _read_odt(self, file_path: str) -> str:
        """Read LibreOffice Writer (.odt) file"""
        if not self.has_odf:
            raise RuntimeError("LibreOffice support not available. Install: pip install odfpy")

        from odf import text, teletype
        from odf.opendocument import load

        doc = load(file_path)
        text_content = []

        # Extract all text elements
        for element in doc.getElementsByType(text.P):
            paragraph_text = teletype.extractText(element)
            if paragraph_text.strip():
                text_content.append(paragraph_text)

        return "\n".join(text_content)

    def _read_ods(self, file_path: str) -> str:
        """Read LibreOffice Calc (.ods) file"""
        if not self.has_odf:
            raise RuntimeError("LibreOffice support not available. Install: pip install odfpy")

        from odf.opendocument import load
        from odf.table import Table, TableRow, TableCell
        from odf import teletype

        doc = load(file_path)
        text_content = []

        # Extract all tables
        for table in doc.spreadsheet.getElementsByType(Table):
            table_name = table.getAttribute('name')
            text_content.append(f"\n=== Sheet: {table_name} ===\n")

            for row in table.getElementsByType(TableRow):
                row_values = []
                for cell in row.getElementsByType(TableCell):
                    cell_text = teletype.extractText(cell)
                    if cell_text:
                        row_values.append(cell_text)

                if row_values:
                    text_content.append(" | ".join(row_values))

        return "\n".join(text_content)

    def _read_text(self, file_path: str) -> str:
        """Read plain text or markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _read_csv(self, file_path: str) -> str:
        """Read CSV file"""
        text_content = []

        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.reader(f)

            for row in csv_reader:
                # Join row values with pipes for better readability
                row_text = " | ".join([cell.strip() for cell in row if cell.strip()])
                if row_text:
                    text_content.append(row_text)

        return "\n".join(text_content)

    def _read_html(self, file_path: str) -> str:
        """Read HTML file and extract text"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback: read raw HTML
            self.log("⚠️  BeautifulSoup not available. Reading raw HTML. Install: pip install beautifulsoup4")
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Break into lines and remove leading/trailing space
        lines = (line.strip() for line in text.splitlines())

        # Drop blank lines
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text

    def _read_with_pandoc(self, file_path: str) -> str:
        """Use pypandoc as universal fallback converter"""
        import pypandoc

        self.log(f"Using pypandoc to convert {Path(file_path).suffix}")

        # Convert to plain text
        text = pypandoc.convert_file(file_path, 'plain')
        return text

    def _read_ipynb(self, file_path: str) -> str:
        """
        Read Jupyter Notebook (.ipynb) file

        Extracts markdown and code cells into readable text format.
        Time Complexity: O(n) where n = number of cells

        Args:
            file_path: Path to .ipynb file

        Returns:
            Formatted text with sections for markdown and code
        """
        try:
            reader = JupyterNotebookReader()
            notebook = reader.read_notebook(file_path)

            text_content = []
            text_content.append("=" * 80)
            text_content.append("JUPYTER NOTEBOOK")
            text_content.append("=" * 80)
            text_content.append("")

            # Get notebook metadata if available
            metadata = notebook.get('metadata', {})
            kernelspec = metadata.get('kernelspec', {})
            if kernelspec:
                kernel_name = kernelspec.get('display_name', kernelspec.get('name', 'Unknown'))
                text_content.append(f"Kernel: {kernel_name}")
                text_content.append("")

            # Process cells - single pass O(n)
            cells = notebook.get('cells', [])
            for i, cell_data in enumerate(cells, 1):
                cell_type = cell_data.get('cell_type', 'unknown')
                source = cell_data.get('source', [])

                # Convert source to text
                if isinstance(source, list):
                    source_text = ''.join(source)
                else:
                    source_text = str(source)

                if not source_text.strip():
                    continue

                # Format based on cell type
                if cell_type == 'markdown':
                    text_content.append(f"[MARKDOWN CELL {i}]")
                    text_content.append("-" * 40)
                    text_content.append(source_text)
                    text_content.append("")

                elif cell_type == 'code':
                    text_content.append(f"[CODE CELL {i}]")
                    text_content.append("-" * 40)
                    text_content.append(source_text)

                    # Include outputs if present
                    outputs = cell_data.get('outputs', [])
                    if outputs:
                        text_content.append("")
                        text_content.append("[OUTPUT]")
                        for output in outputs:
                            output_type = output.get('output_type', '')

                            if output_type == 'stream':
                                stream_text = output.get('text', [])
                                if isinstance(stream_text, list):
                                    stream_text = ''.join(stream_text)
                                text_content.append(stream_text)

                            elif output_type == 'execute_result' or output_type == 'display_data':
                                data = output.get('data', {})
                                # Prefer text/plain representation
                                if 'text/plain' in data:
                                    plain_text = data['text/plain']
                                    if isinstance(plain_text, list):
                                        plain_text = ''.join(plain_text)
                                    text_content.append(plain_text)

                            elif output_type == 'error':
                                error_name = output.get('ename', 'Error')
                                error_value = output.get('evalue', '')
                                text_content.append(f"{error_name}: {error_value}")

                    text_content.append("")

                elif cell_type == 'raw':
                    text_content.append(f"[RAW CELL {i}]")
                    text_content.append("-" * 40)
                    text_content.append(source_text)
                    text_content.append("")

            # Add summary
            summary = reader.get_notebook_summary(notebook)
            text_content.append("=" * 80)
            text_content.append("NOTEBOOK SUMMARY")
            text_content.append("=" * 80)
            text_content.append(f"Total Cells: {summary.get('total_cells', 0)}")
            text_content.append(f"Code Cells: {summary.get('code_cells', 0)}")
            text_content.append(f"Markdown Cells: {summary.get('markdown_cells', 0)}")
            text_content.append(f"Total Code Lines: {summary.get('total_code_lines', 0)}")

            functions = summary.get('functions_defined', [])
            if functions:
                text_content.append(f"Functions Defined: {', '.join(functions)}")

            classes = summary.get('classes_defined', [])
            if classes:
                text_content.append(f"Classes Defined: {', '.join(classes)}")

            text_content.append("=" * 80)

            return "\n".join(text_content)

        except Exception as e:
            raise wrap_exception(
                e,
                DocumentReadError,
                f"Error reading Jupyter notebook: {file_path}",
                context={"file_path": file_path}
            )

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """
        Get dictionary of supported formats based on available libraries

        Returns:
            Dict with categories and supported extensions
        """
        supported = {
            "Always Supported": [".txt", ".md", ".markdown", ".csv", ".ipynb"],
        }

        if self.has_pdf:
            supported["PDF"] = [".pdf"]

        if self.has_docx:
            supported["Microsoft Word"] = [".docx", ".doc"]

        if self.has_openpyxl:
            supported["Microsoft Excel"] = [".xlsx", ".xls"]

        if self.has_odf:
            supported["LibreOffice"] = [".odt", ".ods"]

        try:
            from bs4 import BeautifulSoup
            supported["HTML"] = [".html", ".htm"]
        except ImportError:
            pass

        if self.has_pypandoc:
            supported["Pandoc (Universal)"] = ["Many formats via pypandoc"]

        return supported

    def log(self, message: str):
        """Log message if verbose"""
        if self.verbose:
            print(f"[DocumentReader] {message}")


def main():
    """Test document reader"""
    import argparse

    parser = argparse.ArgumentParser(description="Test document reader")
    parser.add_argument("file", help="File to read")
    args = parser.parse_args()

    reader = DocumentReader(verbose=True)

    print("\n📚 Supported Formats:")
    for category, extensions in reader.get_supported_formats().items():
        print(f"  {category}: {', '.join(extensions)}")

    print(f"\n📄 Reading: {args.file}\n")
    print("=" * 80)

    text = reader.read_document(args.file)
    print(text)

    print("=" * 80)
    print(f"\n✅ Extracted {len(text)} characters")


if __name__ == "__main__":
    main()
